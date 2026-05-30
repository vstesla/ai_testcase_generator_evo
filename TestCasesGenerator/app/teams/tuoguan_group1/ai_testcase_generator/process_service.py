import pandas as pd
import asyncio
import logging
import os
import time
import re
from io import BytesIO
from typing import List, Dict, Any, Tuple
from app.common.cos import ObjectStorage
from app.common.db.db_utils import db_utils
from app.teams.tuoguan_group1.ai_testcase_generator.llm_service import LLMService
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER
from datetime import datetime
import xml.sax.saxutils as saxutils
import zipfile

logger = logging.getLogger(__name__)

# ==================== 常量定义 ====================

# 正则表达式：匹配英文/数字/符号，用于字体替换
REGEX_ENGLISH_NUM = r'([a-zA-Z0-9\-\.\/:,_@]+)'

# 字体模板：用于包裹英文/数字
FONT_TEMPLATE_ENGLISH = r'<font name="TimesNewRoman">\1</font>'

# 字体文件配置：别名 -> 文件名
FONT_FILES = {
    'MicrosoftYaHei': 'msyh.ttc',           # 微软雅黑常规
    'MicrosoftYaHei-Bold': 'msyhbd.ttc',    # 微软雅黑粗体
    'TimesNewRoman': 'times.ttf',           # Times New Roman 常规
    'TimesNewRoman-Bold': 'timesbd.ttf',    # Times New Roman 粗体
    'TimesNewRoman-Italic': 'timesi.ttf',   # Times New Roman 斜体
    'TimesNewRoman-BoldItalic': 'timesbi.ttf',  # Times New Roman 粗斜体
}

# 默认字体名称
FONT_NAME_DEFAULT = 'MicrosoftYaHei'
FONT_NAME_BOLD = 'MicrosoftYaHei-Bold'

# 备用字体（注册失败时使用）
FONT_NAME_FALLBACK = 'Helvetica'
FONT_NAME_BOLD_FALLBACK = 'Helvetica-Bold'

class ProcessService:
    def __init__(self):
        """
        初始化 ProcessService，配置日志、LLM 服务、对象存储及数据库连接池。
        """
        self.object_storage = ObjectStorage()
        self.llm_service = LLMService()
        import asyncio
        self.upload_lock = asyncio.Lock()
        
        # 固定配置
        self.FILE_SOURCE = "tuoguan1"
        self.SKILL_DESCRIPTION = "ai_testcase_generator"
        self.USER_ID = "IT703434"
        self.USER_NAME = "梁金伟"
        # 定义路径配置
        CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
        self.TEMP_DIR = os.path.join(CURRENT_DIR, "Temp")
        self.OUTPUT_DIR = os.path.join(CURRENT_DIR, "Output")
    
    async def _safe_upload_file(self, local_path: str, cos_path: str, user_id: str, user_name: str) -> str:
        """
        线程安全且防并发冲突的文件上传包装方法。
        解决底层 file_id_generator 在并发 SELECT max + 1 时导致的数据库主键冲突问题。
        """
        import asyncio
        async with self.upload_lock:
            return await asyncio.to_thread(
                self.object_storage.upload_file,
                local_path,
                cos_path,
                user_id,
                user_name
            )

    def _register_fonts(self) -> Tuple[str, str]:
        """
        注册所有字体文件，返回常规和粗体字体名称。

        Returns:
            Tuple[str, str]: (常规字体名称, 粗体字体名称)
        """
        font_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")
        print(f"[调试用] 字体目录: {font_dir}")  # 调试用

        font_name = FONT_NAME_DEFAULT
        font_name_bold = FONT_NAME_BOLD

        try:
            # 注册所有字体
            for font_alias, font_file in FONT_FILES.items():
                font_path = os.path.join(font_dir, font_file)
                print(f"[调试用] 注册字体: {font_alias}, 路径: {font_path}, 存在: {os.path.exists(font_path)}")  # 调试用
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont(font_alias, font_path))
                    print(f"[调试用] 字体注册成功: {font_alias}")  # 调试用
                else:
                    print(f"[调试用] 字体文件不存在: {font_path}")  # 调试用

            print(f"[调试用] 所有字体注册完成")  # 调试用
        except Exception as e:
            print(f"[调试用] 字体注册失败: {type(e).__name__}: {e}")  # 调试用
            import traceback
            print(f"[调试用] 字体注册异常堆栈:\n{traceback.format_exc()}")  # 调试用
            logger.warning(f"Font registration failed: {e}")
            font_name = FONT_NAME_FALLBACK
            font_name_bold = FONT_NAME_BOLD_FALLBACK

        return font_name, font_name_bold

    async def process_dataset(self, cos_path: str, file_id: str, file_path_identifier: str, enable_generalization: bool = True, enable_adversarial: bool = False):
        """
        全流程处理：
        1. 下载原始文件 (COS -> Local)
        2. 解析并拆分 -> 3个中间文件 (Local -> COS)
        3. 下载中间文件 -> 读取 (COS -> Local)
        4. 泛化/对抗生成 -> 结果文件 (Local -> COS)
        """
        logger.info(f"Start processing. Original: {cos_path} (ID: {file_id})")
        
        timestamp = int(time.time() * 1000)
        test_case_id = f"TC_{timestamp}"
        
        # 临时目录
        os.makedirs(self.TEMP_DIR, exist_ok=True)
        local_original_path = os.path.join(self.TEMP_DIR, f"download_{os.path.basename(cos_path)}")
        
        try:
            # Step 1: 下载原始文件
            logger.info("Step 1: Downloading original file...")
            # ObjectStorage.download_file 使用 file_id
            if not await asyncio.to_thread(self.object_storage.download_file, file_id, local_original_path):
                 # [模拟下载] 仅用于测试
                if not os.path.exists(local_original_path):
                    # TODO: [DEPLOY_CLEANUP_START]
                    import shutil
                    template_path = os.path.join("Templets", "SmartJudgeConfigs.xlsx")
                    if os.path.exists(template_path):
                        shutil.copy(template_path, local_original_path)
                    else:
                        raise FileNotFoundError("Download failed and no template found")
                    # TODO: [DEPLOY_CLEANUP_END]

            # Step 2: 解析并拆分为3个中间文件，并上传
            logger.info("Step 2: Parsing and splitting to intermediate files...")
            # intermediate_files 现在是 {key: file_id}
            intermediate_files = await self._parse_split_and_upload(local_original_path, file_path_identifier)
            
            # Step 3 & 4 & 5: 读取中间文件 -> 泛化 -> 生成结果 -> 上传
            logger.info("Step 3-5: Processing intermediate files for generalization...")
            
            all_results = []
            if enable_generalization:
                all_results = await self._process_generalization_from_cos(intermediate_files, enable_adversarial)
            else:
                logger.info("Generalization skipped.")

            # Step 6: 保存最终结果并上传
            output_filename = f"{test_case_id}.xlsx"
            output_path = os.path.join(self.OUTPUT_DIR, output_filename)
            os.makedirs(self.OUTPUT_DIR, exist_ok=True)
            
            # 上传最终结果
            file_id, download_url = await self._save_upload_result(all_results, output_path, output_filename, file_path_identifier)

            return {
                "TestCaseID": test_case_id,
                "TestCaseGenStatus": "Y",
                "Message": f"Successfully processed {len(all_results)} items.",
                "DownloadUrl": download_url,
                "Attachments": [{"attachment_id": file_id, "download_url": download_url}]
            }
            
        except Exception as e:
            logger.error(f"Processing failed: {e}")
            return {
                "TestCaseID": test_case_id,
                "TestCaseGenStatus": "N",
                "Message": str(e),
                "DownloadUrl": None
            }
        finally:
            # 清理临时文件
            if os.path.exists(local_original_path):
                try: os.remove(local_original_path) 
                except OSError: pass

    async def _parse_split_and_upload(self, local_path: str, file_path_identifier: str) -> Dict[str, str]:
        """
        解析原始 Excel，拆分为 3 个文件，上传到 COS
        返回: {key: file_id}
        """
        df = pd.read_excel(local_path)
        required_cols = ['TEST_AUDIT_ID', 'ELEMENT_KEY', 'ELEMENT_NAME', 'ELEMENT_VALUE']
        if not all(col in df.columns for col in required_cols):
            raise ValueError(f"Excel missing required columns: {required_cols}")

        # 定义拆分规则
        split_map = {
            "spvSubjectType": "spvSubjectType.xlsx",
            "prdInvestmentScope": "prdInvestmentScope.xlsx",
            "spvInvestmentScope": "spvInvestmentScope.xlsx"
        }
        
        intermediate_file_ids = {}
        
        for key, filename in split_map.items():
            # 筛选数据
            sub_df = df[df['ELEMENT_KEY'] == key]
            if sub_df.empty:
                logger.warning(f"No data found for {key}")
                continue
                
            # 保存到本地
            local_file_path = os.path.join(self.TEMP_DIR, filename)
            sub_df.to_excel(local_file_path, index=False)
            
            # 构造 COS 路径: .../intermediate/filename
            cos_path = f"{self.FILE_SOURCE}/{self.SKILL_DESCRIPTION}/{file_path_identifier}/intermediate/{filename}"
            
            # 上传并获取 ID
            file_id = await self._safe_upload_file(
                local_file_path,
                cos_path,
                self.USER_ID,
                self.USER_NAME
            )
            
            intermediate_file_ids[key] = file_id
            
            # 清理本地
            os.remove(local_file_path)
            
        return intermediate_file_ids

    async def _process_generalization_from_cos(self, intermediate_files: Dict[str, str], enable_adversarial: bool) -> List[Dict]:
        """
        从 COS 下载中间文件，进行泛化处理
        intermediate_files: {key: file_id}
        """
        all_results = []
        
        for key, file_id in intermediate_files.items():
            # 为了文件名，我们需要查一下或者简单构造一个临时名
            # 这里简单构造，反正内容是 excel
            filename = f"{key}_{file_id}.xlsx"
            local_path = os.path.join(self.TEMP_DIR, f"download_{filename}")
            
            # 下载
            try:
                await asyncio.to_thread(self.object_storage.download_file, file_id, local_path)
            except Exception as e:
                logger.error(f"Failed to download intermediate file {file_id}: {e}")
                continue
            
            if not os.path.exists(local_path):
                 logger.error(f"Intermediate file not found after download: {local_path}")
                 continue
                 
            # 读取
            try:
                df = pd.read_excel(local_path)
                records = df.to_dict('records')
                
                # 泛化处理
                for row in records:
                    original_val = row.get('ELEMENT_VALUE')
                    gen_val = original_val # 默认
                    
                    # 调用 LLM
                    try:
                        if key == "spvSubjectType":
                            gen_val = await asyncio.to_thread(self.llm_service.generate_spvSubjectType, original_val)
                        elif key == "prdInvestmentScope":
                            gen_val = await asyncio.to_thread(self.llm_service.generate_prdInvestmentScope, original_val)
                        elif key == "spvInvestmentScope":
                            gen_val = await asyncio.to_thread(self.llm_service.generate_spvInvestmentScope, original_val)
                        
                        # 对抗攻击 (如果开启)
                        if enable_adversarial:
                            gen_val = await asyncio.to_thread(self.llm_service.generate_adversarial_data, gen_val)
                            row['ELEMENT_NAME'] = f"{row.get('ELEMENT_NAME')} (Adversarial)"
                    except Exception as e:
                        logger.error(f"LLM generation failed for {key}: {e}")
                        # 降级处理：保持原始值

                    # 收集结果
                    all_results.append({
                        "TEST_AUDIT_ID": row.get('TEST_AUDIT_ID'),
                        "ELEMENT_KEY": key,
                        "ELEMENT_NAME": row.get('ELEMENT_NAME'),
                        "ORIGINAL_VALUE": original_val,
                        "ELEMENT_VALUE": gen_val
                    })
            except Exception as e:
                logger.error(f"Failed to process file {local_path}: {e}")
            finally:
                # 清理下载的中间文件
                if os.path.exists(local_path):
                    try:
                        os.remove(local_path)
                    except OSError:
                        pass
            
        return all_results

    async def _save_upload_result(self, results: List[Dict], local_path: str, filename: str, file_path_identifier: str) -> tuple:
        """
        保存最终结果，上传 COS，返回 (file_id, 下载链接)
        """
        df = pd.DataFrame(results)
        cols = ["TEST_AUDIT_ID", "ELEMENT_KEY", "ELEMENT_NAME", "ORIGINAL_VALUE", "ELEMENT_VALUE"]
        df = df.reindex(columns=cols)
        
        df.to_excel(local_path, index=False)
        
        cos_path = f"{self.FILE_SOURCE}/{self.SKILL_DESCRIPTION}/{file_path_identifier}/{filename}"
        
        file_id = await self._safe_upload_file(
            local_path,
            cos_path,
            self.USER_ID,
            self.USER_NAME
        )
        
        return file_id, self.object_storage.generate_download_url(file_id)

    async def process_ckxy(self, business_process: str, count: int, file_path_identifier: str, enable_comparison: bool = True) -> Dict[str, Any]:
        """
        处理 ckxy (存款协议) 生成流程
        1. 调用 LLM 获取内容
        2. 生成 PDF
        3. 上传到 COS
        4. 记录到 ai_testcase_generate_record
        """
        logger.info(f"Start processing CKXY. Process: {business_process}, Count: {count}, Comp: {enable_comparison}")
        print(f"[调试用] process_ckxy 开始, business_process: {business_process}, count: {count}")  # 调试用

        # 使用公共方法注册字体
        font_name, font_name_bold = self._register_fonts()

        os.makedirs(self.OUTPUT_DIR, exist_ok=True)
        print(f"[调试用] OUTPUT_DIR 已创建: {self.OUTPUT_DIR}")  # 调试用

        download_urls = []
        attachments = []
        generated_file_info = []

        time_str_batch = datetime.now().strftime("%Y%m%d%H%M%S")
        test_case_id_batch = f"TC_{business_process}_{time_str_batch}"

        for i in range(1, count + 1):
            # 命名规则：TC_业务流程_序号_时间戳
            filename = f"TC_{business_process}_{i}_{time_str_batch}.pdf"
            local_pdf_path = os.path.join(self.OUTPUT_DIR, filename)

            try:
                # 1. 调用 LLM
                content = await asyncio.to_thread(self.llm_service.generate_ckxy, business_process)
                
                # 2. 生成 PDF
                doc = SimpleDocTemplate(local_pdf_path, pagesize=A4,
                                        rightMargin=2*cm, leftMargin=2*cm,
                                        topMargin=2*cm, bottomMargin=2*cm)
                
                styles = getSampleStyleSheet()
                
                # 定义标题样式（使用粗体）
                style_title = ParagraphStyle(
                    'Title_Chinese',
                    parent=styles['Heading1'],
                    fontName=font_name_bold,
                    fontSize=18,
                    leading=22,
                    alignment=TA_CENTER,
                    spaceAfter=20
                )

                style_normal = ParagraphStyle(
                    'Normal_Chinese',
                    parent=styles['Normal'],
                    fontName=font_name,
                    fontSize=12,
                    leading=15,
                    wordWrap='CJK'
                )

                style_normal_bold = ParagraphStyle(
                    'Normal_Chinese_Bold',
                    parent=styles['Normal'],
                    fontName=font_name_bold,
                    fontSize=12,
                    leading=15,
                    wordWrap='CJK'
                )

                # 处理换行与特殊字符
                paragraphs = []
                
                # 插入动态标题：业务流程 + "存款协议"
                pdf_title = f"{business_process}存款协议"
                paragraphs.append(Paragraph(pdf_title, style_title))
                paragraphs.append(Spacer(1, 1*cm))
                
                # 如果内容中混入了字面量 \n，先将其替换为真实的换行符
                content = content.replace('\\n', '\n')
                
                for line in content.split('\n'):
                    # 过滤掉首尾空白，如果为空行则跳过（或者可以加一个更小的 Spacer）
                    line = line.strip()
                    if line:
                        # ReportLab Paragraph 解析类似 XML 的标签，因此需要对 <, >, & 等字符进行转义
                        escaped_line = saxutils.escape(line)

                        # 处理 **加粗** 语法，转换为 font 标签使用粗体字体
                        import re
                        escaped_line = re.sub(r'\*\*(.+?)\*\*', r'<font name="MicrosoftYaHei-Bold">\1</font>', escaped_line)

                        # 英文/数字使用 TimesNewRoman 字体
                        escaped_line = re.sub(r'([a-zA-Z0-9\-\.\/:,_@]+)', r'<font name="TimesNewRoman">\1</font>', escaped_line)

                        p = Paragraph(escaped_line, style_normal)
                        paragraphs.append(p)
                        paragraphs.append(Spacer(1, 0.2*cm))

                # 如果所有内容都为空，添加一个默认提示，防止 build 报错
                if not paragraphs:
                    paragraphs.append(Paragraph("未生成任何有效内容", style_normal))

                doc.build(paragraphs)
                print(f"[调试用] PDF生成成功, local_pdf_path: {local_pdf_path}, 文件大小: {os.path.getsize(local_pdf_path)}")  # 调试用

                # 3. 上传到 COS
                cos_path = f"{self.FILE_SOURCE}/{self.SKILL_DESCRIPTION}/{file_path_identifier}/ckxy/{filename}"
                file_id = await self._safe_upload_file(
                    local_pdf_path,
                    cos_path,
                    self.USER_ID,
                    self.USER_NAME
                )
                
                # 获取下载链接
                download_url = self.object_storage.generate_download_url(file_id)
                download_urls.append(download_url)
                attachments.append({"attachment_id": file_id, "download_url": download_url})
                
                # 存款协议目前没有预期数据，所以 expected_data 为空字典 {}
                generated_file_info.append({
                    "file_id": file_id,
                    "audit_id": test_case_id_batch,
                    "data": {}, 
                    "local_path": local_pdf_path,
                    "download_url": download_url,
                    "filename": filename,
                    "business_process": business_process
                })

            except Exception as e:
                logger.error(f"Failed to generate CKXY {i}: {e}")
            finally:
                # 尽量清理本地生成的 PDF
                if os.path.exists(local_pdf_path):
                    try:
                        os.remove(local_pdf_path)
                    except OSError:
                        pass
        
        if enable_comparison:
            asyncio.create_task(self._parse_and_evaluate(test_case_id_batch, generated_file_info))
            is_comparison_done = False
        else:
            is_comparison_done = True
        
        status = "Y" if download_urls else "N"
        message = f"Successfully generated {len(download_urls)} CKXY PDFs." if download_urls else "Failed to generate CKXY PDFs."
        
        return {
            "TestCaseID": test_case_id_batch,
            "TestCaseGenStatus": status,
            "Message": message,
            "DownloadUrl": ";".join(download_urls),
            "Attachments": attachments,
            "is_comparison_done": is_comparison_done
        }

    def _save_jktzs_intermediate(self, df_list, prefix, timestamp, file_path_identifier):
        """
        保存缴款通知书的中间过程数据（如泛化、对抗后的结果）为 Excel，并上传 COS。
        """
        if not df_list: return
        import pandas as pd
        temp_df = pd.DataFrame(df_list)
        filename = f"{prefix}_jktzs_temp_{timestamp}.xlsx" if prefix else f"jktzs_temp_{timestamp}.xlsx"
        local_path = os.path.join(self.TEMP_DIR, filename)
        temp_df.to_excel(local_path, index=False)
        cos_path = f"{self.FILE_SOURCE}/{self.SKILL_DESCRIPTION}/{file_path_identifier}/intermediate/{filename}"
        import asyncio
        async def _upload_task():
            try:
                await self._safe_upload_file(local_path, cos_path, self.USER_ID, self.USER_NAME)
            except Exception as e:
                logger.error(f"Intermediate upload failed: {e}")
        asyncio.create_task(_upload_task())

    async def _do_jktzs_generalization(self, test_audit_id, raw_data, current_data):
        """
        调用大模型对单笔缴款通知书数据进行泛化生成。

        Args:
            test_audit_id: 测试审计ID
            raw_data: 原始数据字典（用于LLM输入）
            current_data: 当前数据字典（将被更新）
        """
        try:
            import asyncio
            import json
            logger.info(f"[Generalization] Starting for test_audit_id: {test_audit_id}")
            logger.debug(f"[Generalization] Input raw_data: {raw_data}")

            gen_result = await asyncio.to_thread(self.llm_service.generate_jktzs, **raw_data)
            logger.debug(f"[Generalization] LLM returned type: {type(gen_result).__name__}")
            logger.debug(f"[Generalization] LLM returned value: {gen_result}")

            # 如果返回的是字符串，尝试解析为JSON
            if isinstance(gen_result, str):
                try:
                    gen_result = json.loads(gen_result)
                    logger.debug(f"[Generalization] Parsed string to dict: {gen_result}")
                except json.JSONDecodeError as je:
                    logger.warning(f"[Generalization] Failed to parse as JSON: {je}")
                    logger.warning(f"[Generalization] Raw result: {gen_result[:200] if len(gen_result) > 200 else gen_result}")

            # 如果是字典，更新current_data
            if isinstance(gen_result, dict):
                current_data.update(gen_result)
                logger.info(f"[Generalization] Successfully updated current_data for {test_audit_id}")
                logger.debug(f"[Generalization] Updated current_data: {current_data}")
            else:
                logger.warning(f"[Generalization] LLM did not return a dict for {test_audit_id}, type: {type(gen_result).__name__}")

        except Exception as e:
            logger.error(f"[Generalization] Failed for {test_audit_id}: {type(e).__name__}: {e}")
            import traceback
            logger.error(traceback.format_exc())

    async def _do_jktzs_adversarial(self, test_audit_id, current_data):
        """
        调用大模型对单笔缴款通知书数据进行对抗样本生成。

        Args:
            test_audit_id: 测试审计ID
            current_data: 当前数据字典（将被更新，包含泛化后的数据）
        """
        try:
            import asyncio
            import json
            logger.info(f"[Adversarial] Starting for test_audit_id: {test_audit_id}")
            logger.debug(f"[Adversarial] Input current_data: {current_data}")

            adv_result = await asyncio.to_thread(self.llm_service.generate_jktzs_adversarial, **current_data)
            logger.debug(f"[Adversarial] LLM returned type: {type(adv_result).__name__}")
            logger.debug(f"[Adversarial] LLM returned value: {adv_result}")

            # 如果返回的是字符串，尝试解析为JSON
            if isinstance(adv_result, str):
                try:
                    adv_result = json.loads(adv_result)
                    logger.debug(f"[Adversarial] Parsed string to dict: {adv_result}")
                except json.JSONDecodeError as je:
                    logger.warning(f"[Adversarial] Failed to parse as JSON: {je}")
                    logger.warning(f"[Adversarial] Raw result: {adv_result[:200] if len(adv_result) > 200 else adv_result}")

            # 如果是字典，更新current_data
            if isinstance(adv_result, dict):
                current_data.update(adv_result)
                logger.info(f"[Adversarial] Successfully updated current_data for {test_audit_id}")
                logger.debug(f"[Adversarial] Updated current_data: {current_data}")
            else:
                logger.warning(f"[Adversarial] LLM did not return a dict for {test_audit_id}, type: {type(adv_result).__name__}")

        except Exception as e:
            logger.error(f"[Adversarial] Failed for {test_audit_id}: {type(e).__name__}: {e}")
            import traceback
            logger.error(traceback.format_exc())

    def _extend_df_list(self, df_list, group, current_data):
        """
        辅助方法：将新生成的数据字典转换为 DataFrame 并追加到列表中。
        """
        df_list.extend([{**row.to_dict(), 'ELEMENT_VALUE': current_data.get(str(row['ELEMENT_KEY']), row['ELEMENT_VALUE'])} for _, row in group.iterrows()])

    async def _prepare_jktzs_data(self, jktzs_file_path: str, enable_generalization: bool, enable_adversarial: bool, timestamp: int, file_path_identifier: str) -> Dict[str, Any]:
        """
        准备缴款通知书数据：读取 Excel，根据开关决定是否调用 LLM 进行泛化和对抗生成。
        """
        import pandas as pd
        df = pd.read_excel(jktzs_file_path)
        grouped = df.groupby('TEST_AUDIT_ID')
        processed_data_map = {}
        os.makedirs(self.TEMP_DIR, exist_ok=True)
        fh_df_list, dk_df_list, final_df_list = [], [], []
        
        for test_audit_id, group in grouped:
            raw_data = {str(row['ELEMENT_KEY']): row['ELEMENT_VALUE'] for _, row in group.iterrows()}
            current_data = raw_data.copy()
            
            if enable_generalization:
                await self._do_jktzs_generalization(test_audit_id, raw_data, current_data)
                self._extend_df_list(fh_df_list, group, current_data)
                    
            if enable_adversarial:
                await self._do_jktzs_adversarial(test_audit_id, current_data)
                if not enable_generalization:
                    self._extend_df_list(dk_df_list, group, current_data)
                        
            if enable_generalization and enable_adversarial:
                self._extend_df_list(final_df_list, group, current_data)
                
            processed_data_map[test_audit_id] = current_data
            
        if enable_generalization and not enable_adversarial: self._save_jktzs_intermediate(fh_df_list, "FH", timestamp, file_path_identifier)
        elif not enable_generalization and enable_adversarial: self._save_jktzs_intermediate(dk_df_list, "DK", timestamp, file_path_identifier)
        elif enable_generalization and enable_adversarial:
            self._save_jktzs_intermediate(fh_df_list, "FH", timestamp, file_path_identifier)
            self._save_jktzs_intermediate(final_df_list, "", timestamp, file_path_identifier)
            
        return processed_data_map

    def _parse_docx_template(self, jktzs_template_path: str):
        """
        解析 Word (docx) 模板，提取段落文本、表格数据结构及列宽。
        """
        try:
            import docx
            doc_template = docx.Document(jktzs_template_path)
            paragraphs_template = [p.text for p in doc_template.paragraphs if p.text.strip()]
            tables_template, tables_col_widths = [], []
            for table in doc_template.tables:
                table_data = []
                try: col_widths = [col.width if col.width else 1 for col in table.columns]
                except AttributeError: col_widths = None
                for row in table.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
                if table_data:
                    tables_template.append(table_data)
                    tables_col_widths.append(col_widths)
            return paragraphs_template, tables_template, tables_col_widths
        except Exception as e:
            logger.error(f"Failed to read docx template: {e}")
            raise ValueError("缴款通知书模板读取失败，请确保上传了正确的 docx 文件")

    def _calculate_horizontal_spans(self, table_data, style_cmds):
        """
        计算表格的水平合并单元格指令 (SPAN)
        规则1：合并连续的空单元格，用于消除无文本单元格的内部竖线。
        规则2：合并相邻且文本完全一致的单元格，用于还原 Word 模板中原生大段落的合并效果。
        """
        for r_idx, row in enumerate(table_data):
            c_idx = 0
            while c_idx < len(row):
                if row[c_idx].strip() == "":
                    # 规则1: 合并连续的空单元格
                    empty_end = c_idx
                    while empty_end + 1 < len(row) and row[empty_end + 1].strip() == "":
                        empty_end += 1
                    if empty_end > c_idx:
                        style_cmds.append(('SPAN', (c_idx, r_idx), (empty_end, r_idx)))
                    c_idx = empty_end + 1
                else:
                    # 规则2: 合并相邻且文本一致的单元格 (还原 Word 原生合并单元格)
                    text_end = c_idx
                    while text_end + 1 < len(row) and row[text_end + 1] == row[c_idx]:
                        text_end += 1
                    if text_end > c_idx:
                        style_cmds.append(('SPAN', (c_idx, r_idx), (text_end, r_idx)))
                    c_idx = text_end + 1

    def _calculate_vertical_spans(self, table_data, style_cmds):
        """
        计算表格的垂直合并单元格指令 (SPAN)
        规则：按列遍历，向下合并文本完全一致的非空单元格，用于还原 Word 模板的纵向合并效果。
        """
        num_cols = len(table_data[0])
        for c_idx in range(num_cols):
            r_idx = 0
            while r_idx < len(table_data):
                span_end_r = r_idx
                while span_end_r + 1 < len(table_data) and table_data[span_end_r + 1][c_idx] == table_data[r_idx][c_idx]:
                    span_end_r += 1
                if span_end_r > r_idx and table_data[r_idx][c_idx].strip() != "":
                    style_cmds.append(('SPAN', (c_idx, r_idx), (c_idx, span_end_r)))
                r_idx = span_end_r + 1

    def _calculate_table_spans(self, table_data):
        """
        聚合表格的边框、对齐方式以及水平和垂直合并指令 (SPAN)。
        """
        style_cmds = [
            ('GRID', (0, 0), (-1, -1), 0.5, '#000000'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER')
        ]
        self._calculate_horizontal_spans(table_data, style_cmds)
        self._calculate_vertical_spans(table_data, style_cmds)
        return style_cmds

    def _get_pdf_styles(self):
        """
        注册中文字体并获取 ReportLab 的 PDF 段落及表格样式 (包含居中、右对齐等)。
        """
        from reportlab.lib.enums import TA_CENTER, TA_RIGHT
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        # 使用公共方法注册字体
        font_name, font_name_bold = self._register_fonts()

        styles = getSampleStyleSheet()
        style_title = ParagraphStyle('Title_Style', parent=styles['Heading1'], fontName=font_name_bold, fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=15)
        style_normal = ParagraphStyle('Normal_Chinese', parent=styles['Normal'], fontName=font_name, fontSize=10.5, leading=15, firstLineIndent=21, wordWrap='CJK')
        style_bold = ParagraphStyle('Bold_Chinese', parent=styles['Normal'], fontName=font_name_bold, fontSize=10.5, leading=15, wordWrap='CJK')
        style_right = ParagraphStyle('Right_Chinese', parent=styles['Normal'], fontName=font_name, fontSize=10.5, leading=15, alignment=TA_RIGHT, wordWrap='CJK')
        style_table_cell = ParagraphStyle('Table_Cell_Chinese', parent=styles['Normal'], fontName=font_name, fontSize=10.5, leading=12, wordWrap='CJK')
        style_table_cell_bold = ParagraphStyle('Table_Cell_Chinese_Bold', parent=styles['Normal'], fontName=font_name_bold, fontSize=10.5, leading=12, wordWrap='CJK')
        return style_title, style_normal, style_bold, style_right, style_table_cell, style_table_cell_bold


    def _render_single_paragraph(self, p_text, data, i, total_len, styles_tuple):
        """
        渲染单个段落，替换变量，处理中英文混排及高度占位，返回 ReportLab Flowable 列表。
        """
        style_title, style_normal, style_bold, style_right, _, _ = styles_tuple
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.units import cm
        import re
        for key, val in data.items():
            p_text = p_text.replace(f"${key}$", str(val))
            original_key = key.replace('_', '/')
            p_text = p_text.replace(f"${original_key}$", str(val))

        # 英文/数字替换逻辑：通过正则匹配，动态包裹 TimesNewRoman 字体标签
        p_text = re.sub(REGEX_ENGLISH_NUM, FONT_TEMPLATE_ENGLISH, p_text)
        p_text = p_text.replace('\n', '<br/>')
        
        # 修复 XML 标签嵌套可能导致的高度爆炸问题：移除首尾多余的 font 标签或不闭合标签
        try:
            # 简单的测试渲染，如果不报错就继续
            Paragraph(p_text, style_normal)
        except Exception:
            # 如果解析 XML 失败，说明文本中存在特殊闭合或嵌套，回退到纯文本无混合字体模式，防止流程中断
            p_text = re.sub(r'<[^>]+>', '', p_text)

        # 针对前5段的标题（包含"附件"、"确认单"），应用标题样式（已使用粗体字体）
        is_title = (i == 0) or (i < 5 and ("附件" in p_text or "确认单" in p_text))
        if is_title: return [Paragraph(f"<para align='center'>{p_text.strip()}</para>", style_title), Spacer(1, 0.2*cm)]
        # 最后三段的日期/落款部分，应用右对齐样式
        elif i >= total_len - 3 and "附件" not in p_text and "确认单" not in p_text:
            return [Paragraph(p_text, style_right), Spacer(1, 0.2*cm)]
        else:
            return [Paragraph(p_text, style_normal), Spacer(1, 0.2*cm)]

    def _render_single_table(self, table_data, col_widths_ratio, data, style_table_cell, style_table_cell_bold):
        """
        渲染单个表格，替换单元格变量，处理多段落换行、动态列宽及单元格合并。
        """
        from reportlab.platypus import Table, TableStyle, Spacer
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph
        import re
        processed_table_data = []
        for row in table_data:
            processed_row = []
            for cell_text in row:
                for key, val in data.items():
                    cell_text = cell_text.replace(f"${key}$", str(val))
                    original_key = key.replace('_', '/')
                    cell_text = cell_text.replace(f"${original_key}$", str(val))
                cell_flowables = []
                # 【解决 ReportLab 崩溃的核心机制】
                # 将 python-docx 拼在一起的多段落文本 (以 \n 分隔) 拆解成独立的 Paragraph 列表。
                # 否则长段落会在表格 CJK 折行时引发排版引擎死循环，导致 too large on page 高度爆炸。
                for line in cell_text.split('\n'):
                    line_stripped = line.strip()
                    if line_stripped == "":
                        # 保留空行的高度，使用与背景同色的透明标点代替 &nbsp;，防止实体被直接打印
                        line = "<font color='white'>.</font>"
                    else:
                        line = re.sub(REGEX_ENGLISH_NUM, FONT_TEMPLATE_ENGLISH, line)

                    # 针对表格内的特定行文本强制居中并使用粗体
                    if line_stripped.replace(" ", "") in ["重要提示", "认购信息"]:
                        line = f"<para align='center'>{line_stripped}</para>"
                        cell_flowables.append(Paragraph(line, style_table_cell_bold))
                        continue
                        
                    cell_flowables.append(Paragraph(line, style_table_cell))
                    
                # 将拆分好的多段落列表作为整体放入单元格中，ReportLab 会自动垂直排布它们
                processed_row.append(cell_flowables)
            processed_table_data.append(processed_row)
            
        num_cols = len(table_data[0])
        actual_col_widths = [470 / num_cols] * num_cols
        if col_widths_ratio and len(col_widths_ratio) == num_cols:
            total_w = sum(col_widths_ratio)
            if total_w > 0: actual_col_widths = [max(15.0, (w / total_w) * 470) for w in col_widths_ratio]
            
        t = Table(processed_table_data, colWidths=actual_col_widths, rowHeights=None)
        style_cmds = self._calculate_table_spans(table_data)
        t.setStyle(TableStyle(style_cmds))
        return [t, Spacer(1, 0.5*cm)]

    def _render_paragraphs_and_tables(self, data, paragraphs_template, tables_template, tables_col_widths, styles_tuple):
        """
        按模板顺序组装段落和表格，处理特殊的 $Table$ 占位符。
        """
        pdf_paragraphs = []
        table_index = 0
        total_len = len(paragraphs_template)
        for i, p_text in enumerate(paragraphs_template):
            has_table_placeholder = "$Table$" in p_text
            if has_table_placeholder: p_text = p_text.replace("$Table$", "").strip()
            
            if p_text:
                pdf_paragraphs.extend(self._render_single_paragraph(p_text, data, i, total_len, styles_tuple))
            
            if has_table_placeholder and table_index < len(tables_template):
                pdf_paragraphs.extend(self._render_single_table(tables_template[table_index], tables_col_widths[table_index], data, styles_tuple[4], styles_tuple[5]))
                table_index += 1
        return pdf_paragraphs

    async def _render_single_jktzs_pdf(self, idx, test_case_id, data, business_process, time_str_batch, file_path_identifier, paragraphs_template, tables_template, tables_col_widths, styles_tuple, jktzs_template_path):
        """
        根据单笔数据和模板，渲染生成单份缴款通知书 PDF 文件并上传 COS。

        Args:
            idx: 文件序号
            test_case_id: 批次测试用例ID，与ai_testcase_generate_record表保持一致
            data: 填充数据字典
            business_process: 业务流程名称
            time_str_batch: 批次时间戳字符串
            file_path_identifier: 文件路径标识
            paragraphs_template: 段落模板列表
            tables_template: 表格模板列表
            tables_col_widths: 表格列宽列表
            styles_tuple: 样式元组
            jktzs_template_path: 模板文件路径
        """
        from reportlab.platypus import SimpleDocTemplate
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        pdf_filename = f"TC_{business_process}_{idx+1}_{time_str_batch}.pdf"
        local_pdf_path = os.path.join(self.OUTPUT_DIR, pdf_filename)
        doc = SimpleDocTemplate(local_pdf_path, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)

        pdf_paragraphs = self._render_paragraphs_and_tables(data, paragraphs_template, tables_template, tables_col_widths, styles_tuple)
        doc.build(pdf_paragraphs)

        cos_path = f"{self.FILE_SOURCE}/{self.SKILL_DESCRIPTION}/{file_path_identifier}/jktzs/{pdf_filename}"
        file_id = await self._safe_upload_file(local_pdf_path, cos_path, self.USER_ID, self.USER_NAME)
        download_url = self.object_storage.generate_download_url(file_id)

        return {
            "file_id": file_id, "audit_id": test_case_id, "data": data, "local_path": local_pdf_path,
            "download_url": download_url, "filename": pdf_filename, "business_process": business_process,
            "template_filename": os.path.basename(jktzs_template_path)
        }

    async def _render_jktzs_pdfs(self, processed_data_map, paragraphs_template, tables_template, tables_col_widths, business_process, file_count, file_path_identifier, timestamp, jktzs_template_path):
        """
        批量循环处理测试数据，生成指定数量的缴款通知书 PDF 文件。
        """
        os.makedirs(self.OUTPUT_DIR, exist_ok=True)
        download_urls, attachments, generated_file_info = [], [], []
        time_str_batch = datetime.now().strftime("%Y%m%d%H%M%S")
        test_case_id_batch = f"TC_JKTZS_{time_str_batch}"
        
        audit_ids = list(processed_data_map.keys())
        if len(audit_ids) == 0: raise ValueError("没有解析到有效的测试集数据")
        target_audit_ids = [audit_ids[i % len(audit_ids)] for i in range(file_count)]

        styles_tuple = self._get_pdf_styles()

        for idx, source_audit_id in enumerate(target_audit_ids):
            try:
                data = processed_data_map[source_audit_id]
                info = await self._render_single_jktzs_pdf(idx, test_case_id_batch, data, business_process, time_str_batch, file_path_identifier, paragraphs_template, tables_template, tables_col_widths, styles_tuple, jktzs_template_path)
                download_urls.append(info["download_url"])
                attachments.append({"attachment_id": info["file_id"], "download_url": info["download_url"]})
                generated_file_info.append(info)
            except Exception as e:
                logger.error(f"Failed to generate or upload PDF {idx+1} for source_audit_id {source_audit_id}: {e}")
                
        return download_urls, attachments, generated_file_info, test_case_id_batch

    # ==================== 字段名称映射配置 ====================
    # 集中管理 element_key 到 element_name 的映射，便于后续扩展
    ELEMENT_NAME_MAP = {
        "zqjc": "债券简称",
        "zqdm": "债券代码",
        "zqmc": "证券名称",
        "zqdm_1": "证券代码",
        "yjkje": "应缴款金额",
        "jkrq": "缴款日期",
        "skr": "收款人",
        "khx": "开户行",
        "skzh": "收款账户",
        "hkbz": "汇款备注",
        "cpmc": "产品名称",
    }

    async def process_jktzs_attachment(
        self, business_process: str, jktzs_file_path: str, jktzs_template_path: str,
        file_count: int, enable_generalization: bool, enable_adversarial: bool,
        file_path_identifier: str, timestamp: int, enable_comparison: bool = True
    ) -> Dict[str, Any]:
        """
        处理缴款通知书附件生成的总入口函数：
        1. 准备/泛化/对抗生成测试数据
        2. 解析 Word 模板结构
        3. 渲染生成多份 PDF 并上传至 COS
        4. 触发异步解析小助比对任务 (若开启)
        """
        logger.info(f"Start processing JKTZS. Gen: {enable_generalization}, Adv: {enable_adversarial}, Count: {file_count}")
        
        # 1. 准备测试数据 (包含可能的 LLM 泛化/对抗调用)
        processed_data_map = await self._prepare_jktzs_data(jktzs_file_path, enable_generalization, enable_adversarial, timestamp, file_path_identifier)
        
        # 2. 提取模板中的段落和表格骨架
        paragraphs_template, tables_template, tables_col_widths = self._parse_docx_template(jktzs_template_path)
        
        # 3. 循环填充数据，生成多份 PDF 文件并上传 COS
        download_urls, attachments, generated_file_info, test_case_id_batch = await self._render_jktzs_pdfs(
            processed_data_map, paragraphs_template, tables_template, tables_col_widths, business_process, file_count, file_path_identifier, timestamp, jktzs_template_path
        )
        
        # 4. 若开启对比开关，则触发后台异步轮询评测任务
        if enable_comparison:
            import asyncio
            asyncio.create_task(self._parse_and_evaluate(test_case_id_batch, generated_file_info))
            is_comparison_done = False
        else:
            is_comparison_done = True
            
        return {
            "TestCaseID": test_case_id_batch,
            "TestCaseGenStatus": "Y" if download_urls else "N",
            "Message": f"Successfully generated {len(download_urls)} JKTZS PDFs.",
            "DownloadUrl": ";".join(download_urls),
            "Attachments": attachments,
            "is_comparison_done": is_comparison_done,
            
            # 以下为提供给 api.py 使用的内部透传参数
            "download_urls": download_urls,
            "test_case_id_batch": test_case_id_batch,
            "business_process": business_process,
            "generated_file_info": generated_file_info
        }
    def _normalize_amount(self, val):
        """
        格式化金额字符串，去除多余符号并保留两位小数，便于后续比对。
        """
        if not val: return ""
        v = str(val).replace(',', '').replace(' ', '').replace('¥', '').replace('￥', '')
        try: return str(float(v))
        except ValueError: return v

    def _normalize_date(self, val):
        """
        格式化日期字符串为 YYYY-MM-DD 格式，统一日期比对标准。

        注意：只比对年月日，忽略时间部分。
        """
        if not val: return ""
        v = str(val).strip()
        import re
        # 匹配日期部分：支持 YYYY-MM-DD、YYYY/MM/DD、YYYY年MM月DD日 等格式
        date_pattern = r'(\d{4})\s*[-/年\.]\s*(\d{1,2})\s*[-/月\.]\s*(\d{1,2})'
        date_match = re.search(date_pattern, v)
        if not date_match: return v

        # 提取并格式化年月日
        year = date_match.group(1)
        month = f"{int(date_match.group(2)):02d}"
        day = f"{int(date_match.group(3)):02d}"

        # 只返回日期部分，忽略时间
        return f"{year}-{month}-{day}"

    def _get_similarity_calculator(self):
        """
        获取字符串相似度计算函数（优先使用 rapidfuzz，降级使用 difflib）。
        """
        try:
            from rapidfuzz import fuzz
            return fuzz.ratio
        except ImportError:
            try:
                from fuzzywuzzy import fuzz
                return fuzz.ratio
            except ImportError:
                import difflib
                return lambda a, b: difflib.SequenceMatcher(None, str(a), str(b)).ratio() * 100

    def _extract_parsed_values(self, parse_rows):
        """
        递归遍历解析小助返回的 JSON 结构，提取目标字段的解析值。
        """
        import json
        parsed_values = {}
        for row in parse_rows:
            if "element_key" in row and "element_value" in row:
                parsed_values[row["element_key"]] = row["element_value"]
            elif "parse_result" in row:
                try:
                    res_json = json.loads(row["parse_result"])
                    if isinstance(res_json, dict): parsed_values.update(res_json)
                except (ValueError, TypeError): pass
            else: parsed_values.update(row)
        return parsed_values

    async def _call_parse_api(self, filename, download_url, file_type):
        """
        调用远程解析小助 API，提交 PDF 文件并获取处理任务的 fileId。
        """
        import requests
        parse_api_url = "http://tgintellidoc.paasst.cmbchina.cn/api/file-upload-ocr"
        headers = {"channelId": "cntb", "Content-Type": "application/json"}
        payload = {
            "fileName": filename, 
            "ecsUrl": download_url, 
            "dataBuffer": [""], 
            "fileType": file_type,
            "parsePrompt": "", 
            "handlePipeline": "ragAndPos", 
            "isEncrypted": False,
            "parseRange": {
                "startKeyword": "", 
                "endKeyword": "", 
                "startIdx": "1", 
                "endIdx": ""
                }, 
                "sourceId": 202
        }
        import asyncio
        response = await asyncio.to_thread(requests.post, parse_api_url, json=payload, headers=headers)
        if response.status_code != 200: return None
        resp_json = response.json()
        if resp_json.get("returnCode") == "SUC0000":
            return resp_json.get("handleFileId") or resp_json.get("body", {}).get("handleFileId")
        return None

    async def _poll_parse_result(self, remote_db, handle_file_id):
        """
        轮询远程数据库，根据 fileId 获取解析小助的最终解析结果。
        """
        import asyncio
        for _ in range(30):
            await asyncio.sleep(10)
            try:
                rows = remote_db.execute_query("SELECT * FROM tgidoc.file_parse_result WHERE file_id = %s", (handle_file_id,))
                if rows and len(rows) > 0: return rows
            except Exception as qe:
                logger.error(f"Error querying remote DB: {qe}")
        return []

    async def check_ocr_status(self) -> bool:
        """
        检查OCR服务是否开启。
        发送特定请求到 http://ocr-external.paasuat.cmbchina.cn/cv/ocr，
        如果返回 message 包含 "[10005]调用识别模型接口错误"，则说明服务没打开，返回 False。
        否则认为服务已开启，返回 True。
        """
        import httpx
        url = "http://ocr-external.paasuat.cmbchina.cn/cv/ocr"
        payload = {
            "channelCode": "882710489925676032",
            "type": "t2c",
            "imageType": "png",
            "image": "iVBORw0KGgoAAAANSUhEUgAAAIgAAAAkCAIAAADdBuWQAAAACXBIWXMAAA7EAAAOxAGVKw4bAAAJaklEQVRoge2aeUwT3xbHpyxFxEplEUgLSBo2USTBwB+CUQJhUyDumkjQECSBBEJCQBAJxmhRDLIpRCCGRESbALZAgYgIKjWmLEKFIMhaaJGGQksLlM7M++O+N5nXQi3F93v9o5+/bs+cc5c5c793AQKKopAW9Pb2+vr6auNp4K9g9P/ugIGtMSRGTzEkRk8xJEZPMSRGTzEkRk8xJEZPMSRGTzEkRk8xJEZPMent7dXGb2RkxHAl809C0PKuzMA/jEHK9BRDYvSUXSWGw+HIZLLx8fG8vLylpSXNzgqFYmJiQt3+5csXpVK5ZQiTyeTxeCrGurq61dVV3TqsG2KxuLOzU4fAlpYWBEF0a1SXNYbP53d0dLx+/bqrqysxMTEpKenYsWNDQ0Pj4+NCofD69esEAkE9SiaTWVtbFxQUqNizs7Pz8/MTExNV7DAMe3l5+fj41NXVQRC0uLhoa2s7Nzfn4uJSWFiYlJQEQdDq6qqxsbG5uflOhwDo6+tjMpn79+8nEonAgiDI0tKShYUFvs76+vrR0dE3b94EBAQsLCyUl5eTSCQsBIIgqVRqYWFhZPRfX/n8/PyLFy9SUlLu3LmjQ992kJjJyclv376JRCIIgm7fvv358+cjR44YGRkJBAIvLy8wY65cucLn85ubmy0tLVXCFQqFmZmZenNUKpXL5drb26vYi4uL7969m52dbW5uvrGxkZ+f//z5cw6H09HRERcXZ2xsDEFQfX29paUlg8EwMTHRYfAIgkgkEjKZjFn6+/tPnDjR39/v7u6+XdTKyorK6Pz8/Pz8/IqLi1VysytQnbC2thaLxaC8vLxMIpFAWSwWOzg4sNnsLaMgCCpRg0wmCwQCFc/BwcHAwECRSFRbW5uVlQWMvb29ly9fhmH4yZMndDr9j51sbW11cnIKDAzUflzV1dVFRUV4y8bGhuYQuVzu7e0tl8u1b0Ub/"
        }
        try:
            async with httpx.AsyncClient(timeout=10.0) as client:
                response = await client.post(url, json=payload)
                data = response.json()
                message = data.get("message", "")
                if "[10005]调用识别模型接口错误" in message:
                    return False
                return True
        except Exception as e:
            logger.error(f"Failed to check OCR status: {e}")
            return False

    def _calculate_field_score(self, k, exp_val, pars_val, strict_fields, file_type, calc_similarity):
        """
        比对单个字段：根据严格模式或模糊匹配（Levenshtein 距离）计算匹配得分。
        """
        if k in strict_fields:
            norm_exp = self._normalize_amount(exp_val) if k == 'yjkje' else (self._normalize_date(exp_val) if k == 'jkrq' else exp_val)
            norm_pars = self._normalize_amount(pars_val) if k == 'yjkje' else (self._normalize_date(pars_val) if k == 'jkrq' else pars_val)
            if norm_exp == norm_pars: return 100.0, "通过"
            return 0.0, "误识别"
        
        score = calc_similarity(exp_val, pars_val)
        if file_type == "ZLFJ_JKTZSHRGXY":
            if score >= 100.0: return score, "通过"
            if score >= 90.0: return score, "疑似误识别"
            return score, "误识别"
            
        if score >= 90.0: return score, "通过"
        if score >= 80.0: return score, "疑似误识别"
        return score, "误识别"

    def _evaluate_fields(self, expected_data, parsed_values, current_whitelist, strict_fields, fuzzy_fields, file_type, file_sub_type, audit_id, file_id, handle_file_id, db):
        """
        遍历测试数据和解析结果，逐个字段进行比对并汇总单文件的评测明细。

        Args:
            expected_data: 预期数据字典
            parsed_values: 解析结果字典
            current_whitelist: 当前白名单配置
            strict_fields: 严格匹配字段集合
            fuzzy_fields: 模糊匹配字段集合
            file_type: 文件类型编码
            file_sub_type: 文件子类型（Bonds/Securities/空字符串）
            audit_id: 测试用例ID
            file_id: 文件ID
            handle_file_id: 解析任务ID
            db: 数据库连接
        """
        total_valid_fields = 0
        all_pass, any_pass = True, False
        calc_similarity = self._get_similarity_calculator()
        has_whitelist_config = bool(current_whitelist["strict"] or current_whitelist["fuzzy"])

        for k, v in expected_data.items():
            if has_whitelist_config and k not in strict_fields and k not in fuzzy_fields: continue
            exp_val = str(v).strip()
            pars_val = str(parsed_values.get(k, "")).strip()
            if not exp_val: continue

            total_valid_fields += 1
            score, status = self._calculate_field_score(k, exp_val, pars_val, strict_fields, file_type, calc_similarity)

            compare_sql = "INSERT INTO attachments_compare_result (test_case_id, file_id, file_type, file_sub_type, element_key, expected_value, parsed_value, similarity_score, match_status, create_time, update_time, handle_file_id) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW(), %s)"
            db.execute_update(compare_sql, (audit_id, file_id, file_type, file_sub_type, k, exp_val, pars_val, score, status, handle_file_id))
            if status == "通过": any_pass = True
            if status != "通过": all_pass = False

        return total_valid_fields, all_pass, any_pass

    def _get_file_type(self, business_process):
        """
        根据业务流程名称，映射获取解析小助对应的 fileType 参数。
        """
        if "定存" in business_process or "活期" in business_process or "存款协议" in business_process: return "ZLFJ_CKXY"
        file_type_map = {
            "标的合同": "ZLFJ_BDHT", 
            "存款协议": "ZLFJ_CKXY", 
            "成立公告": "ZLFJ_CLGG", 
            "监管备案函": "ZLFJ_JGBAH",
            "缴款通知书": "ZLFJ_JKTZSHRGXY", 
            "认购协议": "ZLFJ_JKTZSHRGXY", 
            "上市开放式基金权益分派预付款通知": "ZLFJ_OPF-QYYFKTZ",
            "托管补充协议": "ZLFJ_TGBCXY", 
            "提取支取函": "ZLFJ_TQZQH", 
            "委托发放现金权益确认书": "ZLFJ_WTFFXJQYQRS", 
            "纸质指令": "ZLFJ_ZZZL"
        }
        return file_type_map.get(business_process, "ZLFJ_JKTZSHRGXY")

    def _get_file_sub_type(self, file_type: str, template_filename: str) -> str:
        """
        根据模板文件名解析出文件的子类型（如证券、债券等）。
        主要用于缴款通知书/认购协议这种在同一个 fileType 下，但不同业务类型需要不同处理的场景。
        """
        if not template_filename or file_type != "ZLFJ_JKTZSHRGXY":
            return ""
        
        # 通过解析文件名中包含的关键字来判断
        if "债券" in template_filename:
            return "Securities"
        elif "证券" in template_filename:
            return "Bonds"
        
        return ""

    def _adjust_whitelist(self, file_type, template_filename, strict_fields, fuzzy_fields):
        """
        根据模板子类型（如证券、债券）动态调整评测的白名单字段，剔除不需要比对的字段。

        Args:
            file_type: 文件类型编码
            template_filename: 模板文件名，可能包含子类型信息
            strict_fields: 严格匹配字段集合（将被修改）
            fuzzy_fields: 模糊匹配字段集合（将被修改）
        """
        # 仅处理缴款通知书/认购协议类型
        if not template_filename or file_type != "ZLFJ_JKTZSHRGXY": return

        # 获取子类型并调整白名单
        file_sub_type = self._get_file_sub_type(file_type, template_filename)

        # 证券类型(Bonds)：移除证券简称和证券代码字段
        if file_sub_type == "Bonds":
            fuzzy_fields.discard('zqjc')
            strict_fields.discard('zqdm')
        # 债券类型(Securities)：移除证券名称和证券代码_1字段
        elif file_sub_type == "Securities":
            fuzzy_fields.discard('zqmc')
            strict_fields.discard('zqdm_1')

    async def _evaluate_single_file(self, info, remote_db, db):
        """
        协调单份生成的 PDF 文件的解析和评测流程（调用 API -> 轮询 DB -> 字段比对）。

        Args:
            info: 文件信息字典，包含file_id、audit_id、data等字段
            remote_db: 远程数据库连接（用于获取解析结果）
            db: 本地数据库连接（用于保存评测结果）
        """
        # 提取文件基本信息
        file_id = info["file_id"]
        audit_id = info["audit_id"]
        expected_data = info["data"]

        # 步骤1: 调用解析API
        file_type = self._get_file_type(info.get("business_process", "缴款通知书"))
        handle_file_id = await self._call_parse_api(info.get("filename", ""), info.get("download_url", ""), file_type)
        if not handle_file_id: return

        # 步骤2: 轮询等待解析完成
        parse_rows = await self._poll_parse_result(remote_db, handle_file_id)
        if not parse_rows: return

        # 步骤3: 提取解析后的字段值
        parsed_values = self._extract_parsed_values(parse_rows)

        EVALUATION_WHITELIST = {
            "ZLFJ_JKTZSHRGXY": { "strict": {'jkrq', 'yjkje', 'skr', 'skzh', 'zqdm_1', 'zqdm'}, "fuzzy": {'khx', 'zqjc', 'zqmc', 'hkbz', 'cpmc'} },
            "ZLFJ_CKXY": { "strict": set(), "fuzzy": set() },
            "ZLFJ_BDHT": { "strict": {'sfhm', 'sfzh', 'zrjg', 'sxtj'}, "fuzzy": {'bdcplx', 'bdcptzfw', 'fkxjtj', 'fklj', 'yt', 'syqzr', 'xjtj'} }
        }
        current_whitelist = EVALUATION_WHITELIST.get(file_type, {"strict": set(), "fuzzy": set()})
        strict_fields = current_whitelist["strict"].copy()
        fuzzy_fields = current_whitelist["fuzzy"].copy()

        # 获取文件子类型并调整白名单
        template_filename = info.get("template_filename", "")
        file_sub_type = self._get_file_sub_type(file_type, template_filename)
        self._adjust_whitelist(file_type, template_filename, strict_fields, fuzzy_fields)

        total_valid_fields, all_pass, any_pass = self._evaluate_fields(expected_data, parsed_values, current_whitelist, strict_fields, fuzzy_fields, file_type, file_sub_type, audit_id, file_id, handle_file_id, db)

        if not expected_data: eval_res = "解析完成(无预期数据)"
        elif total_valid_fields == 0: eval_res = "不通过"
        elif all_pass: eval_res = "完全正确"
        elif any_pass: eval_res = "部分正确"
        else: eval_res = "不通过"
        
        db.execute_update("INSERT INTO ai_evaluation_result (test_case_id, file_id, evaluation_result, create_time, update_time) VALUES (%s, %s, %s, NOW(), NOW())", (audit_id, file_id, eval_res))
        logger.info(f"Successfully evaluated file {file_id}: {eval_res}")

    async def _parse_and_evaluate(self, test_case_id: str, generated_file_info: List[Dict]):
        """
        后台异步任务：批量处理已生成的 PDF，调用解析并比对，最后将评测结果批量写入数据库。
        """
        from app.common.db.mysql_connector import ConnectionPool, MySQLConnector
        from app.common.db.db_utils import DBUtils
        db = DBUtils()
        try:
            remote_pool = ConnectionPool(host="tgidocst.tdsql.dbdns.cmbchina.cn", port=6666, user="app1", password="0s#fzZU1Rq", database="tgidoc", max_connections=5)
            remote_db = MySQLConnector(remote_pool)
        except Exception as e:
            logger.error(f"Failed to connect to remote DB tgidoc: {e}")
            return
            
        for info in generated_file_info:
            try: await self._evaluate_single_file(info, remote_db, db)
            except Exception as e: logger.error(f"Parse and evaluate failed for {info.get('file_id')}: {e}")
            
        try:
            db.execute_update("UPDATE ai_testcase_generate_record SET is_comparison_done = 1 WHERE test_case_id = %s", (test_case_id,))
        except Exception as e:
            logger.error(f"Failed to update is_comparison_done: {e}")

        # 汇总统计 comparison_info 表
        try:
            self._save_comparison_info(test_case_id, db)
        except Exception as e:
            logger.error(f"Failed to save comparison_info: {e}")

    def _build_filename(self, test_case_id: str, file_type: str, file_sub_type: str, extension: str, count: int = None):
        """
        构建文件名（PDF或ZIP）

        Args:
            test_case_id: 测试用例ID
            file_type: 文件类型
            file_sub_type: 文件子类型
            extension: 文件扩展名（如 '.pdf' 或 '.zip'）
            count: 文件数量（仅用于ZIP）

        Returns:
            str: 构建好的文件名
        """
        parts = [test_case_id]

        if file_type:
            parts.append(file_type)
        if file_sub_type:
            parts.append(file_sub_type)
        if count is not None:
            parts.append(str(count))

        return '_'.join(parts) + extension

    def _cleanup_temp_files(self, temp_files: List[str]):
        """
        清理临时文件列表

        Args:
            temp_files: 需要清理的临时文件路径列表
        """
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                    logger.debug(f"Cleaned up temp file: {temp_file}")
            except Exception as e:
                logger.warning(f"Failed to cleanup temp file {temp_file}: {e}")

    async def _download_single_attachment(self, attachment_id: str) -> tuple:
        """
        下载单个附件到临时目录

        Args:
            attachment_id: 附件ID

        Returns:
            tuple: (临时文件路径, 是否成功)
        """
        temp_pdf_path = os.path.join(self.TEMP_DIR, f"temp_download_{attachment_id}.pdf")

        success = await asyncio.to_thread(
            self.object_storage.download_file,
            attachment_id,
            temp_pdf_path
        )

        if success and os.path.exists(temp_pdf_path):
            logger.info(f"Successfully downloaded: {temp_pdf_path}")
            return temp_pdf_path, True

        logger.warning(f"Failed to download attachment: {attachment_id}")
        return None, False

    async def _download_all_attachments(self, attachments: List[Any]) -> tuple:
        """
        批量下载所有附件

        Args:
            attachments: 附件记录列表

        Returns:
            tuple: (下载成功的文件列表, 临时文件列表)
        """
        downloaded_files = []
        temp_files_to_cleanup = []

        for att in attachments:
            att_id = att.get('attachment_id') if isinstance(att, dict) else att[0]
            temp_path, success = await self._download_single_attachment(att_id)

            if success and temp_path:
                downloaded_files.append(temp_path)
                temp_files_to_cleanup.append(temp_path)

        return downloaded_files, temp_files_to_cleanup

    def _create_single_file_stream(self, file_path: str, test_case_id: str, file_type: str, file_sub_type: str) -> tuple:
        """
        创建单个文件流（支持PDF、XLSX等多种格式）

        Args:
            file_path: 文件路径
            test_case_id: 测试用例ID
            file_type: 文件类型
            file_sub_type: 文件子类型

        Returns:
            tuple: (文件流, 文件名, 是否为zip)
        """
        with open(file_path, 'rb') as f:
            file_stream = BytesIO(f.read())

        # 根据实际文件扩展名确定输出文件名
        _, actual_ext = os.path.splitext(file_path)
        if not actual_ext:
            actual_ext = '.bin'  # 默认扩展名

        filename = self._build_filename(test_case_id, file_type, file_sub_type, actual_ext)
        logger.info(f"Returning single file: {filename}")

        return file_stream, filename, False

    def _create_zip_stream(self, files: List[str], test_case_id: str, file_type: str, file_sub_type: str) -> tuple:
        """
        创建ZIP压缩包文件流（支持PDF、XLSX等多种格式）

        Args:
            files: 文件路径列表
            test_case_id: 测试用例ID
            file_type: 文件类型
            file_sub_type: 文件子类型

        Returns:
            tuple: (文件流, 文件名, 是否为zip)
        """
        file_count = len(files)
        zip_filename = self._build_filename(test_case_id, file_type, file_sub_type, '.zip', file_count)

        zip_stream = BytesIO()
        with zipfile.ZipFile(zip_stream, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for idx, file_path in enumerate(files):
                arcname = self._get_archive_name(file_path, test_case_id, idx)
                zipf.write(file_path, arcname)

        zip_stream.seek(0)
        logger.info(f"Returning ZIP archive: {zip_filename} with {file_count} files")

        return zip_stream, zip_filename, True

    def _get_archive_name(self, file_path: str, test_case_id: str, index: int) -> str:
        """
        获取ZIP内的归档文件名（支持PDF、XLSX等多种格式）

        Args:
            file_path: 文件路径
            test_case_id: 测试用例ID
            index: 文件索引

        Returns:
            str: 归档文件名
        """
        arcname = os.path.basename(file_path).replace('temp_download_', '')

        # 如果文件名包含attachment或过长，则使用序号命名
        if 'attachment' in arcname.lower() or len(arcname) > 50:
            # 根据实际文件扩展名确定输出文件名
            _, actual_ext = os.path.splitext(file_path)
            if not actual_ext:
                actual_ext = '.bin'
            return f"{test_case_id}_{index+1}{actual_ext}"

        return arcname
    
    async def download_files_as_stream(self, test_case_id: str, file_type: str = None, file_sub_type: str = None):
        """
        根据test_case_id从COS下载文件，返回文件流供前端下载（支持PDF、XLSX等多种格式）

        Args:
            test_case_id: 测试用例ID（案例集ID）
            file_type: 文件类型（可选，用于命名）
            file_sub_type: 文件子类型（可选，用于命名）

        Returns:
            tuple: (文件流, 文件名, 是否为zip)
                  - 单个文件: (BytesIO流, "文件名.扩展名", False)
                  - 多个文件: (BytesIO流, "压缩包名.zip", True)
        """
        logger.info(f"Starting download for test_case_id: {test_case_id}")

        # 1. 查询附件记录
        from app.common.db.db_utils import DBUtils
        db = DBUtils()
        attachments = db.select(
            table='ai_testcase_generate_attachments',
            columns=['attachment_id', 'download_url'],
            where={'test_case_id': test_case_id}
        )

        if not attachments:
            raise ValueError(f"未找到test_case_id={test_case_id}的附件记录")

        # 2. 准备临时目录
        os.makedirs(self.TEMP_DIR, exist_ok=True)

        # 3. 下载所有附件
        downloaded_files, temp_files = await self._download_all_attachments(attachments)

        if not downloaded_files:
            raise ValueError("未能成功下载任何文件")

        try:
            # 4. 根据文件数量返回相应的文件流
            if len(downloaded_files) == 1:
                return self._create_single_file_stream(downloaded_files[0], test_case_id, file_type, file_sub_type)
            else:
                return self._create_zip_stream(downloaded_files, test_case_id, file_type, file_sub_type)
        finally:
            # 5. 清理临时文件
            self._cleanup_temp_files(temp_files)
 
    def _save_comparison_info(self, test_case_id: str, db):
        """
        汇总统计比对结果并写入 comparison_info 表。

        统计规则：
        - 按 test_case_id + file_type + file_sub_type + element_key 维度分组
        - comparison_count: 该字段的总比对次数
        - correct_count: similarity_score = 100 的次数
        - mistake_count: similarity_score < 90 的次数
        - unclear_count: 90 <= similarity_score < 100 的次数
        - pass_or_not: correct_count == comparison_count 时为 'Y'，否则为 'N'
        """
        # 查询该 test_case_id 下所有比对结果，按 file_type、file_sub_type、element_key 分组统计
        query_sql = """
            SELECT
                test_case_id,
                file_type,
                file_sub_type,
                element_key,
                COUNT(*) AS comparison_count,
                SUM(CASE WHEN similarity_score = 100 THEN 1 ELSE 0 END) AS correct_count,
                SUM(CASE WHEN similarity_score < 90 THEN 1 ELSE 0 END) AS mistake_count,
                SUM(CASE WHEN similarity_score >= 90 AND similarity_score < 100 THEN 1 ELSE 0 END) AS unclear_count
            FROM attachments_compare_result
            WHERE test_case_id = %s
            GROUP BY test_case_id, file_type, file_sub_type, element_key
        """

        try:
            # 使用 execute_query 获取查询结果，返回 List[Dict]
            rows = db.execute_query(query_sql, (test_case_id,))
            logger.info(f"Query returned {len(rows) if rows else 0} rows for test_case_id: {test_case_id}")

            if not rows:
                logger.warning(f"No comparison data found for test_case_id: {test_case_id}")
                return

            # 准备批量插入数据
            for row in rows:
                # execute_query 返回的是 Dict 列表，使用列名访问
                test_case_id_val = row.get('test_case_id')
                file_type = row.get('file_type', '')
                file_sub_type = row.get('file_sub_type') or ""
                element_key = row.get('element_key', '')
                comparison_count = row.get('comparison_count', 0)
                correct_count = row.get('correct_count', 0)
                mistake_count = row.get('mistake_count', 0)
                unclear_count = row.get('unclear_count', 0)

                # 计算正确率
                correct_percentage = round((correct_count / comparison_count) * 100, 2) if comparison_count > 0 else 0.00

                # 获取字段名称（从映射表获取，未知字段使用 key 本身）
                element_name = self.ELEMENT_NAME_MAP.get(element_key, element_key)

                # 判断是否通过：全部正确才算通过
                pass_or_not = "Y" if correct_count == comparison_count else "N"

                row_data = {
                    'test_case_id': test_case_id_val,
                    'file_type': file_type,
                    'file_sub_type': file_sub_type,
                    'element_key': element_key,
                    'element_name': element_name,
                    'comparison_count': str(comparison_count),
                    'correct_count': str(correct_count),
                    'correct_percentage': str(correct_percentage),
                    'mistake_count': str(mistake_count),
                    'unclear_count': str(unclear_count),
                    'pass_or_not': pass_or_not
                }
                print(f"row_data: {row_data} test_case_id: {test_case_id}")
                if db.insert('comparison_info', row_data):
                    logger.info(f"Successfully saved comparison_info record for element_key: {element_key}")
                else:
                    logger.error(f"Failed to batch insert comparison_info for element_key: {element_key}")

        except Exception as e:
            logger.error(f"Failed to save comparison_info for {test_case_id}: {e}")
            import traceback
            logger.error(traceback.format_exc())
            raise
